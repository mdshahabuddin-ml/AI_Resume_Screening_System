from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
import io

router = APIRouter(
    prefix="/export",
    tags=["Resume Export"]
)


class ExportRequest(BaseModel):

    resume_text: str = Field(
        ...,
        min_length=50
    )

    filename: str = "ats_resume"


@router.post("/txt")
def export_txt(request: ExportRequest):

    content = request.resume_text.encode(
        "utf-8"
    )

    file = io.BytesIO(content)

    filename = (
        request.filename
        if request.filename.endswith(".txt")
        else request.filename + ".txt"
    )

    return StreamingResponse(
        file,
        media_type="text/plain",
        headers={
            "Content-Disposition":
            f'attachment; filename="{filename}"'
        }
    )


@router.post("/docx")
def export_docx(request: ExportRequest):

    try:

        from docx import Document

        document = Document()

        for line in request.resume_text.splitlines():

            line = line.strip()

            if not line:
                document.add_paragraph("")
                continue

            # Treat common section headings as headings
            headings = [
                "PROFESSIONAL SUMMARY",
                "TECHNICAL SKILLS",
                "PROFESSIONAL EXPERIENCE",
                "EDUCATION",
                "PROJECTS",
                "CERTIFICATIONS",
                "ACHIEVEMENTS",
                "LANGUAGES"
            ]

            if line.upper() in headings:

                document.add_heading(
                    line,
                    level=1
                )

            else:

                document.add_paragraph(line)

        output = io.BytesIO()

        document.save(output)

        output.seek(0)

        filename = (
            request.filename
            if request.filename.endswith(".docx")
            else request.filename + ".docx"
        )

        return StreamingResponse(
            output,
            media_type=(
                "application/"
                "vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            headers={
                "Content-Disposition":
                f'attachment; filename="{filename}"'
            }
        )

    except Exception as e:

        raise HTTPException(
            status_code=500,
            detail=f"DOCX generation failed: {str(e)}"
        )


@router.post("/pdf")
def export_pdf(request: ExportRequest):

    try:

        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer
        )
        from reportlab.lib.styles import (
            getSampleStyleSheet
        )
        from reportlab.lib.enums import TA_LEFT

        output = io.BytesIO()

        document = SimpleDocTemplate(
            output,
            pagesize=A4,
            rightMargin=40,
            leftMargin=40,
            topMargin=40,
            bottomMargin=40
        )

        styles = getSampleStyleSheet()

        normal_style = styles["Normal"]
        normal_style.fontName = "Helvetica"
        normal_style.fontSize = 9
        normal_style.leading = 12

        heading_style = styles["Heading2"]
        heading_style.fontName = "Helvetica-Bold"
        heading_style.fontSize = 11
        heading_style.leading = 14
        heading_style.alignment = TA_LEFT

        story = []

        headings = [
            "PROFESSIONAL SUMMARY",
            "TECHNICAL SKILLS",
            "PROFESSIONAL EXPERIENCE",
            "EDUCATION",
            "PROJECTS",
            "CERTIFICATIONS",
            "ACHIEVEMENTS",
            "LANGUAGES"
        ]

        for line in request.resume_text.splitlines():

            line = line.strip()

            if not line:
                story.append(
                    Spacer(1, 5)
                )
                continue

            safe_line = (
                line
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )

            if line.upper() in headings:

                story.append(
                    Paragraph(
                        safe_line,
                        heading_style
                    )
                )

            else:

                story.append(
                    Paragraph(
                        safe_line,
                        normal_style
                    )
                )

        document.build(story)

        output.seek(0)

        filename = (
            request.filename
            if request.filename.endswith(".pdf")
            else request.filename + ".pdf"
        )

        return StreamingResponse(
            output,
            media_type="application/pdf",
            headers={
                "Content-Disposition":
                f'attachment; filename="{filename}"'
            }
        )

    except Exception as e:

        raise HTTPException(
            status_code=500,
            detail=f"PDF generation failed: {str(e)}"
        )