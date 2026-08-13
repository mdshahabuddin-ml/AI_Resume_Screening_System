
import io
from backend.schemas.resume_schema import ResumeData


def build_resume_text(resume: ResumeData) -> str:
    lines = []
    c = resume.contact
    lines.append(c.full_name.upper())
    if c.title:
        lines.append(c.title)
    contact_bits = [b for b in [c.location, c.email, c.phone, c.linkedin, c.github, c.portfolio] if b]
    if contact_bits:
        lines.append(" | ".join(contact_bits))
    lines.append("")

    if resume.summary:
        lines += ["PROFESSIONAL SUMMARY", resume.summary, ""]

    if resume.skills:
        lines += ["TECHNICAL SKILLS", ", ".join(resume.skills), ""]

    if resume.education:
        lines.append("EDUCATION")
        for e in resume.education:
            header = f"{e.degree} — {e.institution}"
            if e.location:
                header += f", {e.location}"
            dates = " - ".join([d for d in [e.start_date, e.end_date] if d])
            if dates:
                header += f" ({dates})"
            lines.append(header)
            if e.gpa:
                lines.append(f"GPA: {e.gpa}")
        lines.append("")

    if resume.experience:
        lines.append("EXPERIENCE")
        for exp in resume.experience:
            header = f"{exp.title} — {exp.company}"
            if exp.location:
                header += f", {exp.location}"
            dates = " - ".join([d for d in [exp.start_date, exp.end_date] if d])
            if dates:
                header += f" ({dates})"
            lines.append(header)
            for b in exp.bullets:
                lines.append(f"• {b}")
        lines.append("")

    if resume.projects:
        lines.append("PROJECTS")
        for p in resume.projects:
            title = p.name
            if p.tech_stack:
                title += f" ({', '.join(p.tech_stack)})"
            lines.append(title)
            if p.description:
                lines.append(p.description)
            for b in p.bullets:
                lines.append(f"• {b}")
        lines.append("")

    if resume.certifications:
        lines.append("CERTIFICATIONS")
        for cert in resume.certifications:
            entry = cert.name
            if cert.issuer:
                entry += f" — {cert.issuer}"
            if cert.date:
                entry += f" ({cert.date})"
            lines.append(entry)
        lines.append("")

    if resume.achievements:
        lines.append("ACHIEVEMENTS")
        for a in resume.achievements:
            lines.append(f"• {a}")
        lines.append("")

    if resume.languages:
        lines += ["LANGUAGES", ", ".join(resume.languages), ""]

    return "\n".join(lines).strip() + "\n"


def build_resume_docx(resume: ResumeData) -> bytes:
    """Single-column, no tables, no images — ATS-safe DOCX."""
    import docx
    from docx.shared import Pt

    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    c = resume.contact
    h = doc.add_paragraph()
    run = h.add_run(c.full_name.upper())
    run.bold = True
    run.font.size = Pt(16)

    if c.title:
        doc.add_paragraph(c.title)

    contact_bits = [b for b in [c.location, c.email, c.phone, c.linkedin, c.github, c.portfolio] if b]
    if contact_bits:
        doc.add_paragraph(" | ".join(contact_bits))

    def add_heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12)

    if resume.summary:
        add_heading("PROFESSIONAL SUMMARY")
        doc.add_paragraph(resume.summary)

    if resume.skills:
        add_heading("TECHNICAL SKILLS")
        doc.add_paragraph(", ".join(resume.skills))

    if resume.education:
        add_heading("EDUCATION")
        for e in resume.education:
            header = f"{e.degree} — {e.institution}"
            dates = " - ".join([d for d in [e.start_date, e.end_date] if d])
            if dates:
                header += f" ({dates})"
            doc.add_paragraph(header)

    if resume.experience:
        add_heading("EXPERIENCE")
        for exp in resume.experience:
            header = f"{exp.title} — {exp.company}"
            dates = " - ".join([d for d in [exp.start_date, exp.end_date] if d])
            if dates:
                header += f" ({dates})"
            doc.add_paragraph(header)
            for b in exp.bullets:
                doc.add_paragraph(b, style="List Bullet")

    if resume.projects:
        add_heading("PROJECTS")
        for p in resume.projects:
            title = p.name + (f" ({', '.join(p.tech_stack)})" if p.tech_stack else "")
            doc.add_paragraph(title)
            for b in p.bullets:
                doc.add_paragraph(b, style="List Bullet")

    if resume.certifications:
        add_heading("CERTIFICATIONS")
        for cert in resume.certifications:
            entry = cert.name + (f" — {cert.issuer}" if cert.issuer else "")
            doc.add_paragraph(entry)

    if resume.achievements:
        add_heading("ACHIEVEMENTS")
        for a in resume.achievements:
            doc.add_paragraph(a, style="List Bullet")

    if resume.languages:
        add_heading("LANGUAGES")
        doc.add_paragraph(", ".join(resume.languages))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_resume_pdf(resume: ResumeData) -> bytes:
    """Single-column, text-based (not image-based) PDF — ATS-safe."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                             leftMargin=0.75 * inch, rightMargin=0.75 * inch,
                             topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("Name", parent=styles["Heading1"], fontSize=18, spaceAfter=2)
    heading_style = ParagraphStyle("SectionHeading", parent=styles["Heading2"],
                                    fontSize=12, spaceBefore=10, spaceAfter=4, textColor="black")
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10.5,
                                 leading=14, alignment=TA_LEFT)
    bullet_style = ParagraphStyle("Bullet", parent=body_style, leftIndent=14, bulletIndent=0)

    story = []
    c = resume.contact
    story.append(Paragraph(c.full_name.upper(), name_style))
    if c.title:
        story.append(Paragraph(c.title, body_style))
    contact_bits = [b for b in [c.location, c.email, c.phone, c.linkedin, c.github, c.portfolio] if b]
    if contact_bits:
        story.append(Paragraph(" | ".join(contact_bits), body_style))
    story.append(Spacer(1, 8))

    def section(title, content_paragraphs):
        story.append(Paragraph(title, heading_style))
        story.extend(content_paragraphs)

    if resume.summary:
        section("PROFESSIONAL SUMMARY", [Paragraph(resume.summary, body_style)])

    if resume.skills:
        section("TECHNICAL SKILLS", [Paragraph(", ".join(resume.skills), body_style)])

    if resume.education:
        paras = []
        for e in resume.education:
            header = f"<b>{e.degree}</b> — {e.institution}"
            dates = " - ".join([d for d in [e.start_date, e.end_date] if d])
            if dates:
                header += f" ({dates})"
            paras.append(Paragraph(header, body_style))
        section("EDUCATION", paras)

    if resume.experience:
        paras = []
        for exp in resume.experience:
            header = f"<b>{exp.title}</b> — {exp.company}"
            dates = " - ".join([d for d in [exp.start_date, exp.end_date] if d])
            if dates:
                header += f" ({dates})"
            paras.append(Paragraph(header, body_style))
            for b in exp.bullets:
                paras.append(Paragraph(f"• {b}", bullet_style))
        section("EXPERIENCE", paras)

    if resume.projects:
        paras = []
        for p in resume.projects:
            title = f"<b>{p.name}</b>" + (f" ({', '.join(p.tech_stack)})" if p.tech_stack else "")
            paras.append(Paragraph(title, body_style))
            for b in p.bullets:
                paras.append(Paragraph(f"• {b}", bullet_style))
        section("PROJECTS", paras)

    if resume.certifications:
        paras = [Paragraph(cert.name + (f" — {cert.issuer}" if cert.issuer else ""), body_style)
                 for cert in resume.certifications]
        section("CERTIFICATIONS", paras)

    if resume.achievements:
        paras = [Paragraph(f"• {a}", bullet_style) for a in resume.achievements]
        section("ACHIEVEMENTS", paras)

    if resume.languages:
        section("LANGUAGES", [Paragraph(", ".join(resume.languages), body_style)])

    doc.build(story)
    return buffer.getvalue()